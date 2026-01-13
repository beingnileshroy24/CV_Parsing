from app.services.doc_generator import doc_generator
from app.schemas import CVData
from docx import Document
import io

def test_masking():
    # 1. Create Data with contact info
    data = CVData(
        personal_details={
            "name": "Hidden Figure",
            "email": "hide@me.com",
            "phone": "1234567890",
            "address": "Secret Lair",
            "linkedin": "linkedin.com/in/secret",
            "summary": "Professional summary."
        },
        education=[],
        experience=[],
        projects=[],
        skills=["Ninja Skills"]
    )
    
    # 2. Simulate Masking (clear data)
    data.personal_details.email = None
    data.personal_details.phone = None
    data.personal_details.address = None
    data.personal_details.linkedin = None
    
    # 3. Generate DOCX with paragraph template
    print("Testing Paragraph Template...")
    stream = doc_generator.generate_docx(data, template_style="paragraph")
    doc = Document(stream)
    text = "\n".join([p.text for p in doc.paragraphs])
    
    # print("DEBUG: Full Text:", text)
    
    # 4. Verify Labels are GONE
    assert "Email:" not in text, "Email label failed to hide"
    assert "Phone:" not in text, "Phone label failed to hide"
    assert "Address:" not in text, "Address label failed to hide"
    assert "LinkedIn:" not in text, "LinkedIn label failed to hide"
    assert "Hidden Figure" in text, "Name should remain"
    
    print("[SUCCESS] Paragraph template masking verified.")
    
    # 5. Generate DOCX with tabular template
    print("Testing Tabular Template...")
    stream_tab = doc_generator.generate_docx(data, template_style="tabular")
    doc_tab = Document(stream_tab)
    text_tab = "\n".join([p.text for p in doc_tab.paragraphs])
    
    # print("DEBUG: Full Tabular Text:", text_tab)

    assert "Email:" not in text_tab, "Email label failed to hide in tabular"
    assert "Phone:" not in text_tab, "Phone label failed to hide in tabular"
    assert "Address:" not in text_tab, "Address label failed to hide in tabular"
    
    print("[SUCCESS] Tabular template masking verified.")

if __name__ == "__main__":
    try:
        test_masking()
    except AssertionError as e:
        print(f"[FAILED] {e}")
    except Exception as e:
        print(f"[ERROR] {e}")
