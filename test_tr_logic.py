from docxtpl import DocxTemplate
from docx import Document
import os

def test_row_expansion():
    # 1. Create a minimal template
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    
    # Try the [tr] syntax
    row = table.rows[1].cells
    row[0].text = "[tr for item in items]{{ item.name }}"
    row[1].text = "{{ item.val }}[/tr]"
    
    template_path = "test_tr_template.docx"
    doc.save(template_path)
    
    # 2. Render
    try:
        tpl = DocxTemplate(template_path)
        context = {
            'items': [{'name': 'A', 'val': 1}, {'name': 'B', 'val': 2}]
        }
        tpl.render(context)
        tpl.save("test_tr_output.docx")
        print("[OK] [tr] syntax rendered successfully.")
    except Exception as e:
        print(f"[FAIL] [tr] syntax: {e}")

    # Try the {% tr %} syntax
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    row = table.rows[1].cells
    row[0].text = "{% tr for item in items %}{{ item.name }}"
    row[1].text = "{{ item.val }}{% tr endfor %}"
    
    template_path = "test_tag_template.docx"
    doc.save(template_path)
    
    try:
        tpl = DocxTemplate(template_path)
        tpl.render(context)
        tpl.save("test_tag_output.docx")
        print("[OK] {% tr %} syntax rendered successfully.")
    except Exception as e:
        print(f"[FAIL] {{% tr %}} syntax: {e}")

if __name__ == "__main__":
    test_row_expansion()
