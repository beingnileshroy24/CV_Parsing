from docx import Document
from docx.shared import Pt, RGBColor

def create_paragraph_template():
    doc = Document()
    doc.add_heading('{{ personal_details.name }}', 0)
    
    p = doc.add_paragraph()
    p.add_run('Email: {{ personal_details.email }}').bold = True
    p.add_run(' | Phone: {{ personal_details.phone }}')
    p.add_run('{% if personal_details.address %}\nAddress: {{ personal_details.address }}{% endif %}')
    
    p2 = doc.add_paragraph()
    p2.add_run('{% if personal_details.linkedin %}LinkedIn: {{ personal_details.linkedin }}{% endif %}')
    p2.add_run('{% if personal_details.github %} | GitHub: {{ personal_details.github }}{% endif %}')
    p2.add_run('{% if personal_details.portfolio %} | Portfolio: {{ personal_details.portfolio }}{% endif %}')
    
    doc.add_paragraph('{% if personal_details.summary %}')
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('{{ personal_details.summary }}')
    doc.add_paragraph('{% endif %}')
    
    doc.add_paragraph('{% if skills %}')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('{% for skill in skills %}{{ skill }}{% if not loop.last %}, {% endif %}{% endfor %}')
    doc.add_paragraph('{% endif %}')
    
    doc.add_paragraph('{% if education %}')
    doc.add_heading('Education', level=1)
    doc.add_paragraph('{% for edu in education %}')
    doc.add_paragraph('{{ edu.degree }} from {{ edu.institution }} ({{ edu.year }})').bold = True
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endif %}')
    
    doc.add_paragraph('{% if experience %}')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('{% for exp in experience %}')
    doc.add_paragraph('{{ exp.role }} - {{ exp.company }} ({{ exp.duration }})').bold = True
    doc.add_paragraph('{% for desc in exp.description %}- {{ desc }}\n{% endfor %}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endif %}')
    
    doc.add_paragraph('{% if projects %}')
    doc.add_heading('Projects', level=1)
    doc.add_paragraph('{% for proj in projects %}')
    doc.add_paragraph('{{ proj.name }}').bold = True
    doc.add_paragraph('{{ proj.description }}')
    doc.add_paragraph('Tech: {% for tech in proj.technologies %}{{ tech }}{% if not loop.last %}, {% endif %}{% endfor %}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endif %}')
    
    doc.add_paragraph('{% for section in custom_sections %}')
    doc.add_heading('{{ section.title }}', level=1)
    doc.add_paragraph('{% for item in section.content %}- {{ item }}\n{% endfor %}')
    doc.add_paragraph('{% endfor %}')
    
    doc.save('templates/paragraph_template.docx')
    # Backward compatibility
    doc.save('templates/cv_template.docx')

def create_tabular_template():
    doc = Document()
    doc.add_heading('{{ personal_details.name }}', 0)
    
    # Contact Info
    p = doc.add_paragraph()
    p.add_run('Email: {{ personal_details.email }} | Phone: {{ personal_details.phone }}')
    p.add_run('{% if personal_details.address %}\nAddress: {{ personal_details.address }}{% endif %}')
    
    # Sections using Tables
    # Education
    doc.add_heading('Education', level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'S.No', 'Institution', 'Degree', 'Year'
    row = table.rows[1].cells
    
    row[0].text = '[tr for edu in education][[SNO]]'
    row[1].text = '{{ edu.institution }}'
    row[2].text = '{{ edu.degree }}'
    row[3].text = '{{ edu.year }}[/tr]'
    
    # Experience
    doc.add_heading('Experience', level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'S.No', 'Company & Role', 'Time Period', 'Description'
    row = table.rows[1].cells
    
    row[0].text = '[tr for exp in experience][[SNO]]'
    row[1].text = '{{ exp.company }}\nRole: {{ exp.role }}'
    row[2].text = '{{ exp.duration }}'
    row[3].text = '{{ exp.description_str }}[/tr]'
    
    # Projects
    doc.add_heading('Projects', level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'S.No', 'Project Name', 'Technologies', 'Description'
    row = table.rows[1].cells
    
    row[0].text = '[tr for proj in projects][[SNO]]'
    row[1].text = '{{ proj.name }}'
    row[2].text = '{{ proj.technologies_str }}'
    row[3].text = '{{ proj.description }}[/tr]'
    
    # Skills and Custom Sections
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('{{ skills|join(", ") }}')
    
    doc.add_paragraph('{% for section in custom_sections %}')
    doc.add_heading('{{ section.title }}', level=1)
    doc.add_paragraph('{% for item in section.content %}- {{ item }}\n{% endfor %}')
    doc.add_paragraph('{% endfor %}')
    
    doc.save('templates/tabular_template.docx')

if __name__ == "__main__":
    import os
    if not os.path.exists('templates'):
        os.makedirs('templates')
    create_paragraph_template()
    create_tabular_template()
