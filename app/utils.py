import io
import pypdf
from fastapi import UploadFile

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extracts text from a PDF file content, including hidden hyperlinks."""
    pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
        
        # Extract Links from Annotations
        try:
            if "/Annots" in page:
                for annot in page["/Annots"]:
                    obj = annot.get_object()
                    if "/A" in obj and "/URI" in obj["/A"]:
                        uri = obj["/A"]["/URI"]
                        text += f" [Link: {uri}] "
        except Exception as e:
            # Ignore annotation errors to strictly preserve text extraction
            print(f"Warning: Failed to extract annotations from PDF page: {e}")
            
    return text

def extract_text_from_docx(file_content: bytes) -> str:
    """Extracts text from a DOCX file content."""
    # Note: python-docx can extract text, but for simple extraction we might need to write a temporary file 
    # or handle the zip structure if we want to avoid disk writes. 
    # However, python-docx accepts a file-like object.
    from docx import Document
    
    doc = Document(io.BytesIO(file_content))
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text += "\n" + " | ".join(row_text)
                
    return text

async def extract_text_from_upload_file(file: UploadFile) -> str:
    content = await file.read()
    if file.filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif file.filename.lower().endswith(".docx"):
        return extract_text_from_docx(content)
    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")
