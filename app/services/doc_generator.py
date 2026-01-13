from docxtpl import DocxTemplate, RichText
from app.schemas import CVData
import io
import os
from copy import deepcopy
import re

class DocGenerator:
    def __init__(self):
        self.template_path = "templates/cv_template.docx"

    def generate_docx(self, data: CVData, template_bytes: bytes = None, template_style: str = "paragraph") -> io.BytesIO:
        try:
            from docx import Document
            
            # Load the template as a python-docx Document first
            if template_bytes and len(template_bytes) > 0:
                template_io = io.BytesIO(template_bytes)
                docx = Document(template_io)
            else:
                actual_path = self.template_path
                if template_style == "tabular":
                    actual_path = "templates/tabular_template.docx"
                elif template_style == "tech6":
                    actual_path = "templates/tech6_template.docx"
                elif template_style == "paragraph":
                    actual_path = "templates/paragraph_template.docx"
                
                if not os.path.exists(actual_path):
                    actual_path = "templates/cv_template.docx"
                docx = Document(actual_path)
            
            # Convert to dictionary and PROCESS LINKS
            context = data.model_dump()
            self._process_links(context, docx) # Pass docx if needed for relationship hacking, but RichText handles basic ones
            
            # STAGE 1: MANUAL EXPANSION on the python-docx Document
            self._manually_expand_tables(docx, context)
            
            # STAGE 2: SAVE and LOAD into DocxTemplate for Jinja rendering
            temp_io = io.BytesIO()
            docx.save(temp_io)
            temp_io.seek(0)
            
            doc = DocxTemplate(temp_io)
            doc.render(context)
            
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            return file_stream
        except Exception as e:
            print(f"Template Error Details: {e}")
            raise RuntimeError(f"Error generating DOCX: {e}")

    def _process_links(self, context, docx):
        """
        Recursively scans the context dictionary for URL strings and converts them
        into docxtpl.RichText objects with hyperlinks.
        Also handles specific fields like 'project.url'.
        """
        # Improved Regex to catch linkedin.com/..., github.com/... without protocol
        URL_REGEX = re.compile(r'(https?://[^\s]+|www\.[^\s]+|(?:linkedin\.com|github\.com|gitlab\.com)/[^\s]+)', re.IGNORECASE)
        
        def convert_text(text):
            if not isinstance(text, str): return text
            
            matches = list(URL_REGEX.finditer(text))
            if not matches:
                return text
            
            rt = RichText()
            last_idx = 0
            for m in matches:
                url = m.group(0)
                start, end = m.span()
                
                # Add text before link
                if start > last_idx:
                    rt.add(text[last_idx:start])
                
                # Add link
                href = url if url.startswith('http') else 'http://' + url
                rt.add(url, url_id=href, color='0000FF', underline=True)
                last_idx = end
            
            # Add remaining text
            if last_idx < len(text):
                rt.add(text[last_idx:])
                
            return rt

        # 1. Personal Details specific Links
        pd = context.get('personal_details', {})
        for field in ['linkedin', 'github', 'portfolio']:
            if pd.get(field):
                val = pd[field]
                # Ensure it's treated as a link
                href = val if val.startswith('http') else 'http://' + val
                pd[field] = RichText(val, url_id=href, color='0000FF', underline=True)

        # 2. Recursive scan for other fields
        def recursive_scan(item):
            if isinstance(item, dict):
                for k, v in item.items():
                    if k == 'url' and v and isinstance(v, str):
                        # Explicit URL field in Project
                        href = v if v.startswith('http') else 'http://' + v
                        item[k] = RichText(v, url_id=href, color='0000FF', underline=True)
                    elif isinstance(v, (dict, list)):
                        recursive_scan(v)
                    elif isinstance(v, str):
                        # Attempt detecting links in descriptions/content
                        item[k] = convert_text(v)
            elif isinstance(item, list):
                for idx, i in enumerate(item):
                    if isinstance(i, (dict, list)):
                        recursive_scan(i)
                    elif isinstance(i, str):
                        item[idx] = convert_text(i)
        
        recursive_scan(context)

    def _preprocess_context(self, context):
        """
        Flattens complex lists in the context into strings for easier
        manual template replacement.
        """
        # Process Experience
        for exp in context.get('experience', []):
            desc = exp.get('description', [])
            if isinstance(desc, list):
                # Standard bullet formatting
                exp['description_str'] = "\n".join(f"- {d}" for d in desc)
            else:
                exp['description_str'] = str(desc) if desc else ""
                
        # Process Projects
        for proj in context.get('projects', []):
            tech = proj.get('technologies', [])
            if isinstance(tech, list):
                proj['technologies_str'] = ", ".join(tech)
            else:
                proj['technologies_str'] = str(tech) if tech else ""

    def _manually_expand_tables(self, docx, context):
        """
        Manually expands table rows based on context data using direct DOM manipulation.
        This avoids regex on XML strings and provides stable rendering.
        """
        import re
        from copy import deepcopy
        
        # Pre-process data for simple string replacement
        self._preprocess_context(context)
        
        LOOP_MARKER_RE = re.compile(r'\[tr\s+for\s+(?P<item>\w+)\s+in\s+(?P<list>\w+)\]', re.IGNORECASE)
        
        # Iterate over all tables
        for table in docx.tables:
            rows_to_process = []
            
            # limit scan to avoid errors? no, scan all rows.
            for i, row in enumerate(table.rows):
                # Robust text extraction
                try:
                    row_text = "".join(cell.text for cell in row.cells)
                except Exception:
                    continue
                
                if "[tr" in row_text.lower():
                    match = LOOP_MARKER_RE.search(row_text)
                    if match:
                        rows_to_process.append((i, match.group('item'), match.group('list')))
            
            # Process found loop rows in reverse order
            for i, item_name, list_name in reversed(rows_to_process):
                template_row = table.rows[i]
                data_list = context.get(list_name, [])
                
                w_tbl = table._tbl
                template_tr = template_row._tr
                
                # If no data, remove the template row and continue
                if not data_list:
                    w_tbl.remove(template_tr)
                    continue
                
                # Insert index base - crucial fix for row ordering
                # We need the actual XML index, not the logical row index
                insert_idx = template_tr.getparent().index(template_tr)
                
                for idx, data_item in enumerate(data_list):
                    # Clone the row element
                    new_tr = deepcopy(template_tr)
                    
                    # Iterate through all text nodes in the cloned row
                    for node in new_tr.iter():
                        # looking for w:t tag (text)
                        if node.tag.endswith('}t') and node.text:
                            text = node.text
                            
                            # 1. Clean Loop Markers
                            # Simple string replacement is safer than regex on potentially split text 
                            # if we assume the marker is somewhat intact.
                            # For robustness, we use a liberal regex just for cleanup
                            text = re.sub(r'\[tr\s+for.*?\]', '', text, flags=re.IGNORECASE)
                            text = text.replace('[/tr]', '')
                            
                            # 2. Replace S.No
                            text = text.replace('[[SNO]]', str(idx + 1))
                            text = text.replace('[[ SNO ]]', str(idx + 1))
                            
                            # 3. Replace Variable: {{ item.prop }}
                            # We look for the pattern {{ item_name . prop }}
                            prop_pattern = r'\{\{\s*' + re.escape(item_name) + r'\.(\w+)\s*\}\}'
                            matches = list(re.finditer(prop_pattern, text))
                            
                            # Replace in reverse to avoid index issues
                            for m in reversed(matches):
                                prop_name = m.group(1)
                                replacement_val = str(data_item.get(prop_name, ''))
                                # Replace the specific match span
                                start, end = m.span()
                                text = text[:start] + replacement_val + text[end:]
                            
                            node.text = text
                            
                    # Append the populated row to the table
                    # We use insert on the custom XML element wrapper
                    w_tbl.insert(insert_idx + idx, new_tr)
                
                # Finally, remove the original template row
                w_tbl.remove(template_tr)

doc_generator = DocGenerator()

doc_generator = DocGenerator()
