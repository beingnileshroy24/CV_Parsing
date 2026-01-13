from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tech6_template():
    doc = Document()
    
    # 1. Header
    h1 = doc.add_heading('FORM TECH-6', level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    h2 = doc.add_heading('CURRICULUM VITAE (CV)', level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Key Personnel Table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Row 1: Position
    table.rows[0].cells[0].text = "1. Proposed Position:"
    table.rows[0].cells[1].text = "{{ proposed_position }}"
    
    # Row 2: Firm
    table.rows[1].cells[0].text = "2. Name of Firm:"
    table.rows[1].cells[1].text = "{{ firm_name }}"
    
    # Row 3: Name
    table.rows[2].cells[0].text = "3. Name of Staff:"
    table.rows[2].cells[1].text = "{{ personal_details.name }}"
    
    # Row 4: DOB
    table.rows[3].cells[0].text = "4. Date of Birth:"
    table.rows[3].cells[1].text = "{{ personal_details.date_of_birth }}     Nationality: {{ nationality }}"
    
    # Row 5: Membership
    table.rows[4].cells[0].text = "5. Education:"
    
    # 3. Education Table (Nested in Row 5 or separate? Standards usually separate)
    # Let's clean up. TECH-6 usually lists points.
    
    doc.add_paragraph().add_run('\n5. Education:').bold = True
    
    edu_table = doc.add_table(rows=2, cols=3)
    edu_table.style = 'Table Grid'
    hdr = edu_table.rows[0].cells
    hdr[0].text = "Institution"
    hdr[1].text = "Degree(s) or Diploma(s)"
    hdr[2].text = "Year"
    
    row = edu_table.rows[1].cells
    row[0].text = "[tr for edu in education]{{ edu.institution }}"
    row[1].text = "{{ edu.degree }}"
    row[2].text = "{{ edu.year }}[/tr]"
    
    doc.add_paragraph().add_run('\n6. Membership of Professional Associations:').bold = True
    doc.add_paragraph('{% for mem in memberships %}- {{ mem.society }} ({{ mem.date }})\n{% endfor %}')
    
    doc.add_paragraph().add_run('\n7. Other Training:').bold = True
    doc.add_paragraph('{% for tr in training %}- {{ tr.title }} ({{ tr.start_date }} - {{ tr.end_date }})\n{% endfor %}')
    
    doc.add_paragraph().add_run('\n8. Countries of Work Experience:').bold = True
    doc.add_paragraph('{{ countries_of_work|join(", ") }}')
    
    doc.add_paragraph().add_run('\n9. Languages:').bold = True
    lang_table = doc.add_table(rows=2, cols=2)
    lang_table.style = 'Table Grid'
    lang_table.rows[0].cells[0].text = "Language"
    lang_table.rows[0].cells[1].text = "Proficiency"
    
    l_row = lang_table.rows[1].cells
    l_row[0].text = "[tr for lang in languages]{{ lang.language }}"
    l_row[1].text = "{{ lang.proficiency }}[/tr]"
    
    doc.add_paragraph().add_run('\n10. Employment Record:').bold = True
    doc.add_paragraph('(Starting with present position, list in reverse order every employment held.)')
    
    emp_table = doc.add_table(rows=2, cols=4)
    emp_table.style = 'Table Grid'
    # From: To: Employer: Position:
    e_hdr = emp_table.rows[0].cells
    e_hdr[0].text = "Period"
    e_hdr[1].text = "Employer"
    e_hdr[2].text = "Position"
    e_hdr[3].text = "Summary"

    e_row = emp_table.rows[1].cells
    e_row[0].text = "[tr for exp in experience]{{ exp.duration }}"
    e_row[1].text = "{{ exp.company }}"
    e_row[2].text = "{{ exp.role }}"
    e_row[3].text = "{{ exp.description_str }}[/tr]"
    
    doc.add_paragraph().add_run('\n13. Work Undertaken that Best Illustrates Capability:').bold = True
    
    p = doc.add_paragraph()
    p.add_run('Name of assignment or project: ').bold = True
    p.add_run('{{ representative_project.name }}')
    
    p = doc.add_paragraph()
    p.add_run('Year: ').bold = True
    p.add_run('{{ representative_project.year }}')
    
    p = doc.add_paragraph()
    p.add_run('Location: ').bold = True
    p.add_run('{{ representative_project.location }}')
    
    p = doc.add_paragraph()
    p.add_run('Client: ').bold = True
    p.add_run('{{ representative_project.client }}')
    
    p = doc.add_paragraph()
    p.add_run('Main project features: ').bold = True
    p.add_run('{{ representative_project.main_features }}')
    
    p = doc.add_paragraph()
    p.add_run('Positions held: ').bold = True
    p.add_run('{{ representative_project.positions_held }}')
    
    p = doc.add_paragraph()
    p.add_run('Activities performed: ').bold = True
    p.add_run('{{ representative_project.activities }}')

    doc.save('templates/tech6_template.docx')
    print("Created templates/tech6_template.docx")

if __name__ == "__main__":
    create_tech6_template()
